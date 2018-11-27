# coding = utf-8
import docx
from docx import Document


FilePath = "D:\\py-space\\docx2xlsx\\2017Y_ZhiFa_QB.docx"
doc = Document(FilePath)

print("该文件的段落总数是：" + str(len(doc.paragraphs)))
print("------------------------------")
print("读取前20段的内容：")

i = 1
j = 1
for para in doc.paragraphs:
    if para.text!="":
        if para.text[0].isdigit():
            if int(para.text[0])==j:
                print(para.text[0])
                j = j + 1
    if i >= 20:
        break
    else:
        i = i + 1

