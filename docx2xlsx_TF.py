from docx import Document
import openpyxl


docxFile = "D:\\py-space\\docx2xlsx\\2017Y_ZhiFa_QB-TF.docx"
xlsxFile = "D:\\py-space\\docx2xlsx\\01.xlsx"
doc = Document(docxFile)
wb = openpyxl.load_workbook(xlsxFile)
ws = wb['tf']

print("该文件的段落总数是：" + str(len(doc.paragraphs)))
print("------------------------------")
# print("读取前20段的内容：")

# i = 1
j = 1
for para in doc.paragraphs:
    if para.text.rstrip() != "":
        if para.text[0].isdigit():
            ws.cell(row=j+1, column=1, value=j)
            ws.cell(row=j+1, column=4, value=j)
            index = para.text.find('、')
            question = para.text[index+1:]
            ws.cell(row=j+1, column=2, value=question)
            print('\n第' + str(j) + '题的题目已录入。')
            j = j + 1
        elif para.text[0] == '参':
            index = para.text.find(':')
            tips = para.text.rstrip()[index+1:]
            if tips == "正确":
                answer = "对"
            elif tips == "错误":
                answer = "错"
            else:
                answer = "异常"
            ws.cell(row=j, column=10, value=answer)
            print('答案已录入')
    else:
        print('--------------------')
    # if i >= 30:
    #     break
    # else:
    #     i = i + 1

wb.save(xlsxFile)
wb.close()
# doc.close()
